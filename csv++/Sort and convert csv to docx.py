import os
import pandas as pd
from docx import Document

path = os.getcwd()
# Load the CSV file
file_path = f"{path}/Sorted_Evaly.csv"  # file path
data = pd.read_csv(file_path)

# Sort the data by 'column name'
sorted_data = data.sort_values(by="SN").reset_index(drop=True)

# Create a new Word document
doc = Document()
doc.add_heading("Sorted Product Data", level=1)

# Add a table with the headers and data
table = doc.add_table(rows=1, cols=len(sorted_data.columns))
table.style = "Table Grid"

# Add headers to the Word table
for i, column_name in enumerate(sorted_data.columns):
    table.cell(0, i).text = column_name

# Add the rows from the sorted data
for _, row in sorted_data.iterrows():
    row_cells = table.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = str(cell_value)

# Save the document
output_path = "Sorted_Evaly.docx"
doc.save(output_path)

print("The sorted data has been saved to 'Sorted_Evaly.docx'.")
